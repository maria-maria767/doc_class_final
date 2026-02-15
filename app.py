import sqlite3
import io
#from keras.models import load_model
#import tensorflow as tf
from docx import Document
from flask import Flask, render_template, redirect, request, flash, request, send_from_directory
from flask import request
from werkzeug.exceptions import abort
from werkzeug.utils import secure_filename
from utils import tf_functions as tf
import os
import pprint
from pathlib import Path
import uuid
import numpy as np

app = Flask(__name__)
#===Глобальные переменные
USER_NAME = '111'   # логин пользователя
USER_RIGHTS = 3     # права пользователя
TESTED_DOCS = []    # список протестированных документов по id
ID_OWNER = 0        # id владельца сессии
ALLOWED_EXTENSIONS = {"txt", "docx"}

app.config['SECRET_KEY'] = b'my)secret)key'
UPLOAD_FOLDER = 'requests'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def _get_ext_from_original(filename: str) -> str:
    # "file.DOCX " -> "docx"
    return Path((filename or "").strip()).suffix.lower().lstrip(".")
def _make_safe_filename(original_filename: str) -> str:
    """
    Делает безопасное имя файла для хранения/БД.
    Если имя на кириллице и secure_filename "убил" его (например, оставил 'docx'),
    создаём уникальное имя вида doc_<uuid>.docx
    """
    ext = _get_ext_from_original(original_filename)  # без точки
    ext_with_dot = f".{ext}" if ext else ""

    safe = secure_filename(original_filename)

    # secure_filename может вернуть "docx" (без точки), или пусто, или без расширения
    if not safe or Path(safe).suffix.lower() != ext_with_dot:
        safe = f"doc_{uuid.uuid4().hex}{ext_with_dot}"

    return safe


def _docx_to_text(doc: Document) -> str:
    parts = []

    # Параграфы
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Таблицы (часто текст лежит здесь)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    return "\n".join(parts).strip()


def _extract_text_from_upload(file_storage):
    """
    Возвращает: text_document, size_document, link_document (безопасное имя)
    file_storage — request.files['text_document']
    """
    original_name = file_storage.filename or ""
    ext = _get_ext_from_original(original_name)

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Разрешены только файлы .txt и .docx")

    safe_name = _make_safe_filename(original_name)

    # ВАЖНО: перемотать поток на начало на всякий случай
    try:
        file_storage.stream.seek(0)
    except Exception:
        pass

    data = file_storage.read()
    size_bytes = len(data)

    if size_bytes == 0:
        raise ValueError("Файл пустой или не удалось прочитать содержимое")

    if ext == "txt":
        # Можно заменить encoding при необходимости (например, 'cp1251')
        text = data.decode("utf-8", errors="replace").strip()
        if not text:
            raise ValueError("Текстовый файл не содержит читаемого текста")
        return text, size_bytes, safe_name

    if ext == "docx":
        doc = Document(io.BytesIO(data))
        text = _docx_to_text(doc)
        if not text:
            raise ValueError("Не удалось извлечь текст из .docx (возможно, документ пустой)")
        return text, size_bytes, safe_name

    raise ValueError("Неподдерживаемый формат")#===Получение соединения с БД
def get_db_connection():
    conn = sqlite3.connect('confidentiality.db')
    conn.row_factory = sqlite3.Row
    return conn

#===Стартовая страница
@app.route('/')
def index():
    global USER_RIGHTS
    USER_RIGHTS = 3
    return redirect("/auth")


# ===Авторизация пользователей в системе
@app.route('/auth', methods=('GET', 'POST'))
def auth():
    if request.method == 'POST':
        global USER_NAME
        global USER_RIGHTS
        try:
            login_employee = str(request.form['selected_login'])
            USER_NAME = login_employee
            if login_employee == 'korutov':
                USER_RIGHTS = 1
            else:
                USER_RIGHTS = 2
            password_employee = str(request.form['password_employee'])
        except ValueError:
            flash('Введены некорректные значения')
        conn = get_db_connection()
        item = conn.execute("""SELECT * FROM employees
                            WHERE (employees.login = ?)and(employees.password = ?)
             """, (login_employee, password_employee)).fetchall()
        conn.close()
        if len(item) == 0:
            flash('Введен неправильный логин и(или) пароль пользователя! Повторите ввод')
            return redirect(f'/auth')
        else:
            global ID_OWNER
            ID_OWNER = item[0]['id_employee']
            return redirect(f'/correspondents')

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM employees
        """).fetchall()
    conn.close()
    return render_template('auth.html', logins=pos)


def replace_text(paragraph, key, value):
    """ Работа docx - заполнение параграфов """

    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


def replace_text_in_tables(table, key, value):
    """ Работа docx - заполнение таблиц """

    for row in table.rows:
        for cell in row.cells:
            if key in cell.text:
                cell.text = cell.text.replace(key, value)







###############
###############
############### Сотрудники университета
###############
###############

@app.route('/employees')
#===Получение данных из таблицы 'Сотрудники'
def employees():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM employees, departments, positions, role_users
                WHERE (employees.id_department = departments.id_department)AND(employees.id_position = positions.id_position)AND
                (employees.id_role = role_users.id_role)
            """).fetchall()
    conn.close()
    return render_template('employees.html', employees=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного сотрудника из БД
def get_employee(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT employees.id_employee, employees.name, employees.login, employees.password,
                                  employees.id_department, departments.name_department,
                                  employees.email, employees.id_position, positions.position,
                                  employees.id_role, role_users.name_role 
                           FROM employees, departments, positions, role_users
                           WHERE (employees.id_employee = ?)AND(employees.id_department=departments.id_department)AND
                           (employees.id_position=positions.id_position)AND(employees.id_role=role_users.id_role)
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


#===Вывод карточки о сотруднике
@app.route('/employees/<int:id_employee>')
def employee(id_employee):
    pos = get_employee(id_employee)
    return render_template('employee.html', employee=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Добавление нового сотрудника
@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():

    if request.method == 'POST':
        # добавление нового сотрудника в БД после заполнения формы
        try:
            name = request.form['name']
            login = request.form['login']
            password = request.form['password']
            id_department = request.form['selected_id_department']
            email = request.form['email']
            id_position = request.form['selected_id_position']
            id_role = request.form['selected_id_role']

            base_flag = 1
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name and login and password and id_department and email and id_position and id_role):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'employees' ('name', 'login', 'password', 'id_department', 'email', id_position, id_role)  VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (name, login, password, id_department, email, id_position, id_role))
                conn.commit()
                conn.close()
                return redirect('/employees')

    # отрисовка формы
    conn = get_db_connection()
    posd = conn.execute("""SELECT * FROM departments""").fetchall()
    posp = conn.execute("""SELECT * FROM positions""").fetchall()
    posr = conn.execute("""SELECT * FROM role_users""").fetchall()
    conn.close()
    return render_template('new_employee.html', departments=posd, positions=posp, roles=posr, user_name=USER_NAME, user_rights=USER_RIGHTS)


@app.route('/employee/<int:id_employee>/del', methods=('GET', 'POST'))
#===Удаление сотрудника
def del_employee(id_employee):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM employees WHERE id_employee = ?", (id_employee,))
        conn.commit()
        conn.close()
        return redirect('/employees')

    # отрисовка формы
    pos = get_employee(id_employee)
    return render_template('del_employee.html', employee=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/employee/<int:id_employee>/update', methods=('GET', 'POST'))
#===Редактирование сотрудника
def edit_employee(id_employee):
    base_flag = 1
    if request.method == 'POST':
        try:
            name = request.form['name']
            login = request.form['login']
            password = request.form['password']
            id_department = request.form['selected_id_department']
            email = request.form['email']
            id_position = request.form['selected_id_position']
            id_role = request.form['selected_id_role']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name and login and password and id_department and email and id_position and id_role):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'employees' set 'name'=?, 'login'=?, 'password'=?, 'id_department'=?,'email'=?,'id_position'=?, 'id_role'=?  WHERE id_employee = ?",
                    (name, login, password, id_department, email, id_position, id_role, id_employee))
                conn.commit()
                conn.close()
                return redirect('/employees')
    # отрисовка формы
    pos = get_employee(id_employee)
    conn = get_db_connection()
    posd = conn.execute("""SELECT * FROM departments""").fetchall()
    posp = conn.execute("""SELECT * FROM positions""").fetchall()
    posr = conn.execute("""SELECT * FROM role_users""").fetchall()
    conn.close()
    #for pp in pos:
    #    tt=pp.birthday

    return render_template('edit_employee.html', employee=pos, departments=posd, positions=posp, roles=posr, user_name=USER_NAME, user_rights=USER_RIGHTS)

###############
###############
############### Подразделения университета
###############
###############

@app.route('/departments')

#===Получение данных из таблицы 'Подразделения'
def departments():

    conn = get_db_connection()
    pos = conn.execute("""SELECT  * FROM departments
            """).fetchall()
    conn.close()
    return render_template('departments.html', departments=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного подразделения из БД
def get_department(item_id):

    conn = get_db_connection()
    item = conn.execute("""SELECT DISTINCT * FROM departments
                WHERE departments.id_department= ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/departments/<int:id_department>')
#===Вывод карточки о подразделении
def department(id_department):

    pos = get_department(id_department)
    return render_template('department.html', department=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/new_department', methods=('GET', 'POST'))
#===Добавление нового подразделения
def new_department():

    base_flag = 1
    if request.method == 'POST':
        # добавление нового подразделения в БД после заполнения формы
        try:
            name_department = request.form['name_department']
            submission = request.form['submission']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_department and submission):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'departments' ('name_department', 'submission')  VALUES (?, ?)",
                    (name_department, submission))
                conn.commit()
                conn.close()
                return redirect('/departments')

    # отрисовка формы
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM departments""").fetchall()
    conn.close()
    return render_template('new_department.html', departments=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/department/<int:id_department>/del', methods=('GET', 'POST'))
#===Удаление подразделения
def del_department(id_department):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM departments WHERE id_department = ?", (id_department,))
        conn.commit()
        conn.close()
        return redirect('/departments')

    # отрисовка формы
    pos = get_department(id_department)
    return render_template('del_department.html', department=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/department/<int:id_department>/update', methods=('GET', 'POST'))
#===Редактирование подразделения
def edit_department(id_department):
    base_flag = 1
    if request.method == 'POST':
        try:
            name_department = request.form['name_department']
            submission = request.form['submission']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_department and submission):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'departments' set 'name_department'=?, 'submission'=? WHERE id_department = ?",
                    (name_department, submission, id_department))
                conn.commit()
                conn.close()
                return redirect('/departments')
    # отрисовка формы
    pos = get_department(id_department)
    return render_template('edit_department.html', department=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)


###############
###############
############### Роли пользователей
###############
###############

@app.route('/role_users')

#===Получение данных из таблицы 'Роли пользователей'
def role_users():

    conn = get_db_connection()
    pos = conn.execute("""SELECT  * FROM role_users
            """).fetchall()
    conn.close()
    return render_template('role_users.html', role_users=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одной роли пользователей из БД
def get_role_user(item_id):

    conn = get_db_connection()
    item = conn.execute("""SELECT DISTINCT * FROM role_users
                WHERE role_users.id_role= ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/role_users/<int:id_role>')
#===Вывод карточки о роли пользователей
def role_user(id_role):

    pos = get_role_user(id_role)
    return render_template('role_user.html', role_user=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/new_role_user', methods=('GET', 'POST'))
#===Добавление новой роли пользователей
def new_role_user():

    base_flag = 1
    if request.method == 'POST':
        # добавление новой роли пользователей в БД после заполнения формы
        try:
            name_role = request.form['name_role']
            description_role = request.form['description_role']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_role and description_role):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'role_users' ('name_role', 'description_role')  VALUES (?, ?)",
                    (name_role, description_role))
                conn.commit()
                conn.close()
                return redirect('/role_users')

    # отрисовка формы
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM role_users""").fetchall()
    conn.close()
    return render_template('new_role_user.html', role_users=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/role_user/<int:id_role>/del', methods=('GET', 'POST'))
#===Удаление роли пользователей
def del_role_user(id_role):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM role_users WHERE id_role = ?", (id_role,))
        conn.commit()
        conn.close()
        return redirect('/role_users')

    # отрисовка формы
    pos = get_role_user(id_role)
    return render_template('del_role_user.html', role_user=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/role_user/<int:id_role>/update', methods=('GET', 'POST'))
#===Редактирование роли пользователей
def edit_role_user(id_role):
    base_flag = 1
    if request.method == 'POST':
        try:
            name_role = request.form['name_role']
            description_role = request.form['description_role']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_role and description_role):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'role_users' set 'name_role'=?, 'description_role'=? WHERE id_role = ?",
                    (name_role, description_role, id_role))
                conn.commit()
                conn.close()
                return redirect('/role_users')
    # отрисовка формы
    pos = get_role_user(id_role)
    return render_template('edit_role_user.html', role_user=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

###############
###############
############### Типы прав
###############
###############

@app.route('/access_types')

#===Получение данных из таблицы 'Типы прав'
def access_types():

    conn = get_db_connection()
    pos = conn.execute("""SELECT  * FROM access_types
            """).fetchall()
    conn.close()
    return render_template('access_types.html', access_types=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного типа прав из БД
def get_access_type(item_id):

    conn = get_db_connection()
    item = conn.execute("""SELECT DISTINCT * FROM access_types
                WHERE access_types.id_access_type= ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/access_types/<int:id_access_type>')
#===Вывод карточки о типе прав
def access_type(id_access_type):

    pos = get_access_type(id_access_type)
    return render_template('access_type.html', access_type=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/new_access_type', methods=('GET', 'POST'))
#===Добавление нового типа прав
def new_access_type():

    base_flag = 1
    if request.method == 'POST':
        # добавление нового типа прав в БД после заполнения формы
        try:
            name_access_right = request.form['name_access_right']
            description_rights = request.form['description_rights']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_access_right and description_rights):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'access_types' ('name_access_right', 'description_rights')  VALUES (?, ?)",
                    (name_access_right, description_rights))
                conn.commit()
                conn.close()
                return redirect('/access_types')

    # отрисовка формы
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM access_types""").fetchall()
    conn.close()
    return render_template('new_access_type.html', access_types=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/access_type/<int:id_access_type>/del', methods=('GET', 'POST'))
#===Удаление типа прав
def del_access_type(id_access_type):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM access_types WHERE id_access_type = ?", (id_access_type,))
        conn.commit()
        conn.close()
        return redirect('/access_types')

    # отрисовка формы
    pos = get_access_type(id_access_type)
    return render_template('del_access_type.html', access_type=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/access_type/<int:id_access_type>/update', methods=('GET', 'POST'))
#===Редактирование типа прав
def edit_access_type(id_access_type):
    base_flag = 1
    if request.method == 'POST':
        try:
            name_access_right = request.form['name_access_right']
            description_rights = request.form['description_rights']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_access_right and description_rights):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'access_types' set 'name_access_right'=?, 'description_rights'=? WHERE id_access_type = ?",
                    (name_access_right, description_rights, id_access_type))
                conn.commit()
                conn.close()
                return redirect('/access_types')
    # отрисовка формы
    pos = get_access_type(id_access_type)
    return render_template('edit_access_type.html', access_type=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

###############
###############
############### Корреспонденты
###############
###############

@app.route('/correspondents')
#===Получение данных из таблицы 'Корреспонденты'
def correspondents():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM correspondents, type_correspondents
                WHERE (correspondents.id_type_correspondent = type_correspondents.id_type_correspondent)
            """).fetchall()
    conn.close()
    return render_template('correspondents.html', correspondents=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного корреспондента из БД
def get_correspondent(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT correspondents.id_correspondent, correspondents.name_correspondent,
                                  correspondents.id_type_correspondent, type_correspondents.type_correspondent                                   
                           FROM correspondents, type_correspondents
                           WHERE (correspondents.id_correspondent = ?)AND(correspondents.id_type_correspondent=type_correspondents.id_type_correspondent)
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


#===Вывод карточки о корреспонденте
@app.route('/correspondents/<int:id_correspondent>')
def correspondent(id_correspondent):
    pos = get_correspondent(id_correspondent)
    return render_template('correspondent.html', correspondent=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Добавление нового корреспондента
@app.route('/new_correspondent', methods=('GET', 'POST'))
def new_correspondent():

    if request.method == 'POST':
        # добавление нового корреспондента в БД после заполнения формы
        try:
            name_correspondent = request.form['name_correspondent']
            id_type_correspondent = request.form['selected_id_type_correspondent']

            base_flag = 1
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_correspondent and  id_type_correspondent):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'correspondents' ('name_correspondent', 'id_type_correspondent')  VALUES (?, ?)",
                    (name_correspondent, id_type_correspondent))
                conn.commit()
                conn.close()
                return redirect('/correspondents')

    # отрисовка формы
    conn = get_db_connection()
    post = conn.execute("""SELECT * FROM type_correspondents""").fetchall()
    conn.close()
    return render_template('new_correspondent.html', type_correspondents=post, user_name=USER_NAME, user_rights=USER_RIGHTS)


@app.route('/correspondent/<int:id_correspondent>/del', methods=('GET', 'POST'))
#===Удаление корреспондента
def del_correspondent(id_correspondent):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM correspondents WHERE id_correspondent = ?", (id_correspondent,))
        conn.commit()
        conn.close()
        return redirect('/correspondents')

    # отрисовка формы
    pos = get_correspondent(id_correspondent)
    return render_template('del_correspondent.html', correspondent=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/correspondent/<int:id_correspondent>/update', methods=('GET', 'POST'))
#===Редактирование корреспондента
def edit_correspondent(id_correspondent):
    base_flag = 1
    if request.method == 'POST':
        try:
            name_correspondent = request.form['name_correspondent']
            id_type_correspondent = request.form['selected_id_type_correspondent']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_correspondent and id_type_correspondent):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'correspondents' set 'name_correspondent'=?, 'id_type_correspondent'=?  WHERE id_correspondent = ?",
                    (name_correspondent, id_type_correspondent, id_correspondent))
                conn.commit()
                conn.close()
                return redirect('/correspondents')
    # отрисовка формы
    pos = get_correspondent(id_correspondent)
    conn = get_db_connection()
    post = conn.execute("""SELECT * FROM type_correspondents""").fetchall()
    conn.close()
    #for pp in pos:
    #    tt=pp.birthday

    return render_template('edit_correspondent.html', correspondent=pos, type_correspondents=post, user_name=USER_NAME, user_rights=USER_RIGHTS)




###############
###############
############### Документы
###############
###############

@app.route('/documents')
#===Получение данных из таблицы 'Документы'
def documents():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM documents, category_documents, correspondents
                WHERE (documents.id_category_document = category_documents.id_category_document)AND
                (documents.id_correspondent = correspondents.id_correspondent)
            """).fetchall()
    conn.close()
    return render_template('documents.html', documents=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного документа из БД
def get_document(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT documents.id_document, documents.name_document, documents.text_document, documents.size_document,
                                  documents.link_document,                                  
                                  documents.id_category_document, category_documents.category_document,
                                  documents.id_correspondent, correspondents.name_correspondent                                   
                           FROM documents, category_documents, correspondents
                           WHERE (documents.id_document = ?)AND(documents.id_category_document=category_documents.id_category_document)AND
                                 (documents.id_correspondent=correspondents.id_correspondent)
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


#===Вывод карточки о документе
@app.route('/documents/<int:id_document>')
def document(id_document):
    pos = get_document(id_document)
    return render_template('document.html', document=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Добавление нового документа



@app.route('/new_document', methods=('GET', 'POST'))
def new_document():
    if request.method == 'POST':
        try:
            name_document = request.form.get('name_document', '').strip()
            id_category_document = request.form.get('selected_id_category_document')
            id_correspondent = request.form.get('selected_id_correspondent')

            if not name_document:
                flash('Не указано наименование документа')
                return redirect('/new_document')

            if not id_category_document or not id_correspondent:
                flash('Не выбраны категория документа и/или корреспондент')
                return redirect('/new_document')

            if 'text_document' not in request.files:
                flash('Файл не выбран')
                return redirect('/new_document')

            file = request.files['text_document']
            if not file or not file.filename:
                flash('Файл не выбран')
                return redirect('/new_document')

            # Извлекаем текст и размер. Ссылку/имя файла НЕ используем.
            text_document, size_document, _ = _extract_text_from_upload(file)

            # Ссылки на документ больше нет
            link_document = ""

        except ValueError as e:
            flash(str(e))
            return redirect('/new_document')
        except Exception as e:
            flash(f'Ошибка при обработке документа: {e}')
            return redirect('/new_document')

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO documents (name_document, text_document, size_document, link_document, id_category_document, id_correspondent) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (name_document, text_document, size_document, link_document, id_category_document, id_correspondent)
        )
        conn.commit()
        conn.close()

        flash('Документ успешно добавлен')
        return redirect('/documents')

    # GET: отрисовка формы
    conn = get_db_connection()
    postcd = conn.execute(
        "SELECT * FROM category_documents WHERE id_category_document IN (1, 2)"
    ).fetchall()
    postc = conn.execute("SELECT * FROM correspondents").fetchall()
    conn.close()
    return render_template(
        'new_document.html',
        category_documents=postcd,
        correspondents=postc,
        user_name=USER_NAME,
        user_rights=USER_RIGHTS
    )


@app.route('/document/<int:id_document>/del', methods=('GET', 'POST'))
#===Удаление документа
def del_document(id_document):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM documents WHERE id_document = ?", (id_document,))
        conn.commit()
        conn.close()
        return redirect('/documents')

    # отрисовка формы
    pos = get_document(id_document)
    tt=pos['link_document']
    return render_template('del_document.html', document=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/document/<int:id_document>/check', methods=('GET', 'POST'))
#===Изменение категории документа
def check_document(id_document):
    if request.method == 'POST':
        global model
        #model = load_model('workmodel.h5')
        # Required for model to work
        global graph
        #graph = tf.get_default_graph()

        doc = get_document(id_document)
        name_document = doc['name_document']
        text_document = doc['text_document']
        size_document = doc['size_document']
        link_document = doc['link_document']
        id_correspondent = doc['id_correspondent']
        #try:
        #    with open('output.txt', 'r', encoding='utf-8') as f:
        #        content = f.read()
        #except FileNotFoundError:
        #    content = "Файл не найден"

        global TESTED_DOCS
        if (id_document in TESTED_DOCS):
            flash('Для данного документа рекомендуемая категория уже была определена!')
        else:
            #===определение категории документа
            model = tf.load_model("best_model.keras")
            doc_tensor = tf.tensor_create(text_document)
            results = tf.predict(model,doc_tensor)
            r = np.array(results, dtype="float32").reshape(-1)

            if r.size == 1:
                # sigmoid: r[0] = P(class=1)
                p1 = float(r[0])
                p0 = 1.0 - p1
            elif r.size >= 2:
                # softmax: [P(class=0), P(class=1)]
                p0 = float(r[0])
                p1 = float(r[1])
            else:
                raise ValueError(f"Unexpected model output: {results}")


            res_id_category = int(p1 > p0)


            LABELS = {
                0: "Не конфиденциальный",
                1: "Конфиденциальный",
            }
            category_text = LABELS[res_id_category]

            flash("Результаты анализа документа:")
            flash(f"Не конфиденциальный: {p0:.3f}; Конфиденциальный: {p1:.3f}.")
            flash("Выбрана категория: " + category_text)

            TESTED_DOCS.append(id_document)

            DB_CATEGORY_ID = {0: 2, 1: 1}
            id_category_document = DB_CATEGORY_ID[res_id_category]
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE 'documents' set 'name_document'=?, 'text_document'=?, 'size_document'=?, 'link_document'=?, 'id_category_document'=?,'id_correspondent'=?  WHERE id_document = ?",
                (name_document, text_document, size_document, link_document, id_category_document, id_correspondent,
                id_document))
            conn.commit()
            conn.close()
        return redirect('/documents')

    # отрисовка формы
    pos = get_document(id_document)
    tt=pos['link_document']
    return render_template('check_document.html', document=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)



@app.route('/document/<int:id_document>/update', methods=('GET', 'POST'))
#===Редактирование локумента
def edit_document(id_document):
    base_flag = 1
    if request.method == 'POST':
        try:
            name_document = request.form['name_document']
            text_document = request.form['text_document']
            size_document = request.form['size_document']
          #  link_document = request.form['link_document']
            id_category_document = request.form['selected_id_category_document']
            id_correspondent = request.form['selected_id_correspondent']
            link_document = ""
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (name_document and text_document and size_document):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'documents' set 'name_document'=?, 'text_document'=?, 'size_document'=?, 'link_document'=?, 'id_category_document'=?,'id_correspondent'=?  WHERE id_document = ?",
                    (name_document, text_document, size_document, link_document, id_category_document, id_correspondent, id_document))
                conn.commit()
                conn.close()
                return redirect('/documents')
    # отрисовка формы
    pos = get_document(id_document)
    conn = get_db_connection()
    poscd = conn.execute(
    "SELECT * FROM category_documents WHERE id_category_document IN (1, 2)"
).fetchall()
    posc = conn.execute("""SELECT * FROM correspondents""").fetchall()
    conn.close()

    return render_template('edit_document.html', document=pos, category_documents=poscd, correspondents=posc, user_name=USER_NAME, user_rights=USER_RIGHTS)



###############
###############
############### Назначение прав доступа
###############
###############

@app.route('/assignment_rights')
#===Получение данных из таблицы 'Назначение прав доступа'
def assignment_rights():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM assignment_rights, role_users, access_types, documents
                WHERE (assignment_rights.id_role = role_users.id_role)AND(assignment_rights.id_access_type = access_types.id_access_type)AND
                      (assignment_rights.id_document=documents.id_document)
            """).fetchall()
    conn.close()
    return render_template('assignment_rights.html', assignment_rights=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Получение одного назначения прав доступа из БД
def get_assignment_right(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT assignment_rights.id_assignment_rights, assignment_rights.date_assignment,
                                  assignment_rights.id_role, role_users.name_role,
                                  assignment_rights.id_access_type, access_types.name_access_right,
                                  assignment_rights.id_document, documents.name_document 
                           FROM assignment_rights, role_users, access_types, documents
                           WHERE (assignment_rights.id_assignment_rights = ?)AND
                                 (assignment_rights.id_role=role_users.id_role)AND
                                 (assignment_rights.id_access_type=access_types.id_access_type)AND
                                 (assignment_rights.id_document=documents.id_document)
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


#===Вывод карточки о назначении прав
@app.route('/assignment_rights/<int:id_assignment_rights>')
def assignment_right(id_assignment_rights):
    pos = get_assignment_right(id_assignment_rights)
    return render_template('assignment_right.html', assignment_right=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

#===Добавление нового назначения прав
@app.route('/new_assignment_right', methods=('GET', 'POST'))
def new_assignment_right():

    if request.method == 'POST':
        # добавление нового назначения прав в БД после заполнения формы
        try:
            date_assignment = request.form['date_assignment']
            id_role = request.form['selected_id_role']
            id_access_type = request.form['selected_id_access_type']
            id_document = request.form['selected_id_document']

            base_flag = 1
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (date_assignment):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'assignment_rights' ('date_assignment', 'id_role', 'id_access_type', 'id_document')  VALUES (?, ?, ?, ?)",
                    (date_assignment, id_role, id_access_type, id_document))
                conn.commit()
                conn.close()
                return redirect('/assignment_rights')

    # отрисовка формы
    conn = get_db_connection()
    posru = conn.execute("""SELECT * FROM role_users""").fetchall()
    posat = conn.execute("""SELECT * FROM access_types""").fetchall()
    posd = conn.execute("""SELECT * FROM documents""").fetchall()
    conn.close()
    return render_template('new_assignment_right.html', role_users=posru, access_types=posat, documents=posd, user_name=USER_NAME, user_rights=USER_RIGHTS)


@app.route('/assignment_right/<int:id_assignment_rights>/del', methods=('GET', 'POST'))
#===Удаление назначения прав
def del_assignment_right(id_assignment_rights):
    if request.method == 'POST':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM assignment_rights WHERE id_assignment_rights = ?", (id_assignment_rights,))
        conn.commit()
        conn.close()
        return redirect('/assignment_rights')

    # отрисовка формы
    pos = get_assignment_right(id_assignment_rights)
    return render_template('del_assignment_right.html', assignment_right=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

@app.route('/assignment_right/<int:id_assignment_rights>/update', methods=('GET', 'POST'))
#===Редактирование назначения прав
def edit_assignment_right(id_assignment_rights):
    base_flag = 1
    if request.method == 'POST':
        try:
            date_assignment = request.form['date_assignment']
            id_role = request.form['selected_id_role']
            id_access_type = request.form['selected_id_access_type']
            id_document = request.form['selected_id_document']
        except ValueError:
            flash('Некорректные значения')
            base_flag = 0
        if not base_flag > 0:
            flash('Не все поля заполнены')
        else:
            if not (date_assignment):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE 'assignment_rights' set 'date_assignment'=?, 'id_role'=?, 'id_access_type'=?, 'id_document'=?  WHERE id_assignment_rights = ?",
                    (date_assignment, id_role, id_access_type, id_document, id_assignment_rights))
                conn.commit()
                conn.close()
                return redirect('/assignment_rights')
    # отрисовка формы
    pos = get_assignment_right(id_assignment_rights)
    conn = get_db_connection()
    posru = conn.execute("""SELECT * FROM role_users""").fetchall()
    posat = conn.execute("""SELECT * FROM access_types""").fetchall()
    posd = conn.execute("""SELECT * FROM documents""").fetchall()
    conn.close()

    return render_template('edit_assignment_right.html', assignment_right=pos, role_users=posru, access_types=posat, documents=posd, user_name=USER_NAME, user_rights=USER_RIGHTS)



#
# =========Выдача сообщения об отсутствии страницы (404)
#

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

#===Функция выхода из приложения
@app.route('/shutdown', methods=['GET', 'POST'])
def shutdown():
    global USER_NAME
    global USER_RIGHTS
    global TESTED_DOCS
    global ID_OWNER
    USER_NAME = ''  # логин пользователя
    USER_RIGHTS = 3  # права пользователя
    TESTED_DOCS = []  # список протестированных документов по id
    ID_OWNER = 0      # id владельца сессии
    return redirect('/auth')

#===Функция вывода профиля пользователя
@app.route('/profile', methods=['GET', 'POST'])
def profile():
    global ID_OWNER
    global USER_NAME
    global USER_RIGHTS
    pos = get_employee(ID_OWNER)
    return render_template('employee.html', employee=pos, user_name=USER_NAME, user_rights=USER_RIGHTS)

if __name__ == '__main__':
    app.run()